import os
from typing import Optional, BinaryIO
from datetime import datetime
from fastapi import UploadFile, HTTPException
import aiofiles
import magic
import fitz  # PyMuPDF for PDF processing
from docx import Document as DocxDocument  # python-docx for DOCX processing
from openpyxl import load_workbook  # openpyxl for XLSX processing
from PIL import Image  # Pillow for image processing

from app.models.document import Document, DocumentType, DocumentStatus
from app.core.config import settings
from app.db.session import SessionLocal

MIME_TYPE_MAP = {
    'application/pdf': DocumentType.PDF,
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
    'image/jpeg': DocumentType.IMAGE,
    'image/png': DocumentType.IMAGE,
}

class DocumentService:
    def __init__(self):
        self.upload_dir = settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)

    async def save_upload_file(self, file: UploadFile) -> str:
        """Save the uploaded file and return its path."""
        file_path = os.path.join(self.upload_dir, file.filename)
        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await file.read()
            await out_file.write(content)
        return file_path

    def validate_file_type(self, file_path: str) -> DocumentType:
        """Validate file type using magic numbers."""
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        if file_type not in MIME_TYPE_MAP:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        return MIME_TYPE_MAP[file_type]

    async def extract_metadata(self, file_path: str, doc_type: DocumentType) -> dict:
        """Extract metadata based on file type."""
        metadata = {}
        
        try:
            if doc_type == DocumentType.PDF:
                pdf = fitz.open(file_path)
                metadata['page_count'] = len(pdf)
                metadata['has_images'] = any(page.get_images() for page in pdf)
                
                # Extract text to count words and check for tables
                text = ""
                for page in pdf:
                    text += page.get_text()
                metadata['word_count'] = len(text.split())
                # Simple table detection (can be improved)
                metadata['has_tables'] = '|' in text or '\t' in text
                
            elif doc_type == DocumentType.DOCX:
                doc = DocxDocument(file_path)
                metadata['page_count'] = len(doc.sections)
                metadata['word_count'] = len(doc.paragraphs)
                metadata['has_tables'] = len(doc.tables) > 0
                metadata['has_images'] = len(doc.inline_shapes) > 0
                
            elif doc_type == DocumentType.XLSX:
                wb = load_workbook(file_path)
                metadata['page_count'] = len(wb.sheetnames)
                metadata['has_tables'] = True  # Excel files are essentially tables
                
            elif doc_type == DocumentType.IMAGE:
                with Image.open(file_path) as img:
                    metadata['page_count'] = 1
                    metadata['has_images'] = True
                    metadata['word_count'] = 0  # Requires OCR for text detection
                    metadata['has_tables'] = False  # Requires ML for table detection
        
        except Exception as e:
            print(f"Error extracting metadata: {str(e)}")
            # Return basic metadata if extraction fails
            metadata = {
                'page_count': 1,
                'word_count': 0,
                'has_tables': False,
                'has_images': False
            }
        
        return metadata

    async def create_document(
        self,
        file: UploadFile,
        user_id: int,
        db: SessionLocal
    ) -> Document:
        """Create a new document with metadata."""
        # Save file
        file_path = await self.save_upload_file(file)
        
        # Validate and get file type
        doc_type = self.validate_file_type(file_path)
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Create document record
        document = Document(
            name=file.filename,
            type=doc_type,
            size=file_size,
            status=DocumentStatus.PENDING,
            url=f"/uploads/{file.filename}",  # You might want to use a cloud storage URL
            user_id=user_id
        )
        
        db.add(document)
        db.commit()
        
        # Extract metadata asynchronously
        metadata = await self.extract_metadata(file_path, doc_type)
        
        # Update document with metadata
        document.page_count = metadata.get('page_count')
        document.word_count = metadata.get('word_count')
        document.has_tables = metadata.get('has_tables')
        document.has_images = metadata.get('has_images')
        document.status = DocumentStatus.COMPLETED
        document.processed_at = datetime.utcnow()
        
        db.commit()
        db.refresh(document)
        
        return document

    def get_document(self, document_id: int, user_id: int, db: SessionLocal) -> Optional[Document]:
        """Get a document by ID."""
        return db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id
        ).first()

    def get_user_documents(self, user_id: int, db: SessionLocal) -> list[Document]:
        """Get all documents for a user."""
        return db.query(Document).filter(Document.user_id == user_id).all()

    def delete_document(self, document_id: int, user_id: int, db: SessionLocal) -> bool:
        """Delete a document."""
        document = self.get_document(document_id, user_id, db)
        if not document:
            return False
            
        # Delete file
        file_path = os.path.join(self.upload_dir, os.path.basename(document.url))
        if os.path.exists(file_path):
            os.remove(file_path)
            
        db.delete(document)
        db.commit()
        return True 