"""Word document table detection implementation."""
from typing import List, Dict, Any, Optional, Tuple
import logging
from pathlib import Path
import docx
from docx.table import Table as DocxTable, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import numpy as np
from datetime import datetime

from app.services.ml.base import BaseTableDetector
from app.schemas.analysis import (
    TableDetectionParameters,
    TableDetectionResult,
    PageTableInfo,
    DetectedTable,
    TableCell,
    BoundingBox
)

logger = logging.getLogger(__name__)

class WordTableDetector(BaseTableDetector):
    """Service for detecting and extracting tables from Word documents with high accuracy."""
    
    def __init__(self):
        """Initialize the Word table detection service."""
        logger.debug("Initializing Word Table Detection Service")
        super().__init__()
    
    @property
    def supported_formats(self) -> List[str]:
        """List of supported file formats."""
        return ["docx", "doc"]

    async def detect_tables(
        self,
        file_path: str,
        parameters: Dict[str, Any]
    ) -> TableDetectionResult:
        """
        Detect and extract tables from a Word document.
        
        Args:
            file_path: Path to the Word document
            parameters: Dictionary of detection parameters
            
        Returns:
            TableDetectionResult containing detailed table information
            
        Raises:
            ValueError: If file is invalid or processing fails
        """
        logger.info(f"Detecting tables in Word document: {file_path}")
        
        try:
            # Convert dictionary parameters to TableDetectionParameters
            detection_params = TableDetectionParameters(**parameters)
            
            # Validate file before processing
            if not self.validate_file(file_path):
                raise ValueError(f"Invalid Word document: {file_path}")

            # Load document
            doc = docx.Document(file_path)
            
            # Initialize result containers
            pages: List[PageTableInfo] = []
            total_confidence = 0.0
            total_tables = 0
            
            # Process document sections and tables
            current_page = 1
            tables_in_page: List[DetectedTable] = []
            
            for element in doc.element.body:
                if element.tag.endswith('sectPr'):
                    # New section - store current page tables and reset
                    if tables_in_page:
                        pages.append(self._create_page_info(current_page, tables_in_page))
                        tables_in_page = []
                    current_page += 1
                    
                elif element.tag.endswith('tbl'):
                    table = DocxTable(element, doc)
                    detected_table = self._process_table(table, detection_params)
                    
                    if detected_table:
                        tables_in_page.append(detected_table)
                        total_confidence += detected_table.confidence_score
                        total_tables += 1
            
            # Add remaining tables
            if tables_in_page:
                pages.append(self._create_page_info(current_page, tables_in_page))
            
            # Calculate average confidence
            avg_confidence = total_confidence / total_tables if total_tables > 0 else 1.0
            
            # Create and return final result
            return TableDetectionResult(
                pages=pages,
                total_tables=total_tables,
                average_confidence=avg_confidence,
                processing_metadata={
                    "processor": "WordTableDetector",
                    "processed_at": datetime.utcnow().isoformat(),
                    "parameters_used": parameters  # Use original parameters dict
                }
            )
            
        except Exception as e:
            logger.error(f"Table detection failed: {str(e)}", exc_info=True)
            raise ValueError(f"Table detection failed: {str(e)}")

    def validate_file(self, file_path: str) -> bool:
        """
        Validate if the file is a valid Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            bool: True if file is valid, False otherwise
        """
        try:
            if not Path(file_path).exists():
                logger.error(f"File not found: {file_path}")
                return False
                
            docx.Document(file_path)
            return True
            
        except Exception as e:
            logger.error(f"File validation failed: {str(e)}")
            return False

    def _process_table(
        self,
        table: DocxTable,
        parameters: TableDetectionParameters
    ) -> Optional[DetectedTable]:
        """
        Process a single table and extract its content and structure.
        
        Args:
            table: Word document table
            parameters: Table detection parameters
            
        Returns:
            DetectedTable if valid table is found, None otherwise
        """
        try:
            rows = len(table.rows)
            cols = len(table.columns)
            
            # Skip tables that don't meet minimum row requirement
            if rows < parameters.min_row_count:
                return None
            
            # Extract cells with their properties
            cells: List[TableCell] = []
            header_rows: List[int] = []
            has_headers = False
            
            for i in range(rows):
                for j in range(cols):
                    cell = table.cell(i, j)
                    
                    # Skip merged cells that were already processed
                    if self._is_merged_continuation(table, i, j):
                        continue
                    
                    cell_content = self._extract_cell_content(cell)
                    is_header = self._is_header_cell(cell, i == 0)
                    
                    if is_header and i == 0:
                        has_headers = True
                        if i not in header_rows:
                            header_rows.append(i)
                    
                    spans = self._get_cell_spans(table, i, j)
                    
                    cells.append(TableCell(
                        content=cell_content["text"],
                        row_index=i,
                        col_index=j,
                        row_span=spans["row_span"],
                        col_span=spans["col_span"],
                        is_header=is_header,
                        confidence=1.0  # Word tables have explicit structure
                    ))
            
            # Get table dimensions and position
            bbox = self._calculate_table_bbox(table)
            
            return DetectedTable(
                bbox=bbox,
                confidence_score=1.0,
                rows=rows,
                columns=cols,
                cells=cells,
                has_headers=has_headers,
                header_rows=header_rows
            )
            
        except Exception as e:
            logger.error(f"Failed to process table: {str(e)}")
            return None

    def _extract_cell_content(self, cell: _Cell) -> Dict[str, Any]:
        """
        Extract text and formatting from a table cell.
        
        Args:
            cell: Word document table cell
            
        Returns:
            Dictionary containing cell content and formatting
        """
        try:
            text_parts = []
            
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Get vertical alignment safely
            v_align = None
            if cell._tc.tcPr is not None:
                v_align_elements = cell._tc.tcPr.xpath('.//w:vAlign')
                if v_align_elements and hasattr(v_align_elements[0], 'val'):
                    v_align = str(v_align_elements[0].val)
            
            return {
                "text": "\n".join(text_parts),
                "formatting": {
                    "alignment": str(cell.paragraphs[0].alignment if cell.paragraphs else None),
                    "vertical_alignment": v_align
                }
            }
        except Exception as e:
            logger.warning(f"Could not extract cell content completely: {str(e)}")
            return {
                "text": "",
                "formatting": {
                    "alignment": None,
                    "vertical_alignment": None
                }
            }

    def _is_header_cell(self, cell: _Cell, is_first_row: bool) -> bool:
        """
        Determine if a cell is a header cell based on formatting.
        
        Args:
            cell: Word document table cell
            is_first_row: Whether the cell is in the first row
            
        Returns:
            bool: True if cell is a header, False otherwise
        """
        try:
            if not cell.paragraphs:
                return False
                
            paragraph = cell.paragraphs[0]
            
            # Check for header indicators
            indicators = [
                any(run.bold for run in paragraph.runs),
                any(run.style and "head" in run.style.name.lower() for run in paragraph.runs if hasattr(run, 'style') and run.style),
                paragraph.style and "head" in paragraph.style.name.lower()
            ]
            
            # Check vertical alignment for first row
            if is_first_row and cell._tc.tcPr is not None:
                v_align_elements = cell._tc.tcPr.xpath('.//w:vAlign')
                if v_align_elements and hasattr(v_align_elements[0], 'val'):
                    indicators.append(v_align_elements[0].val in ['center', 'top'])
            
            return any(indicators)
            
        except Exception as e:
            logger.warning(f"Could not check header cell formatting: {str(e)}")
            return False

    def _get_cell_spans(self, table: DocxTable, row: int, col: int) -> Dict[str, int]:
        """
        Get row and column spans for a cell.
        
        Args:
            table: Word document table
            row: Row index
            col: Column index
            
        Returns:
            Dictionary containing row_span and col_span
        """
        try:
            cell = table.cell(row, col)
            tc = cell._tc
            
            # Get grid span (column span)
            grid_span = 1
            if tc.tcPr is not None:
                gridSpan = tc.tcPr.xpath('./w:gridSpan')
                if gridSpan and hasattr(gridSpan[0], 'val'):
                    try:
                        grid_span = int(gridSpan[0].val)
                    except (AttributeError, ValueError):
                        grid_span = 1
            
            # Calculate row span
            row_span = 1
            if tc.tcPr is not None:
                vMerge = tc.tcPr.xpath('./w:vMerge')
                if vMerge:
                    # Check if this is the start of a vertical merge
                    is_merge_start = (
                        hasattr(vMerge[0], 'val') and 
                        vMerge[0].val == 'restart'
                    )
                    
                    if is_merge_start or not hasattr(vMerge[0], 'val'):
                        # Count continued merged cells
                        for i in range(row + 1, len(table.rows)):
                            next_cell = table.cell(i, col)
                            if next_cell._tc.tcPr is not None:
                                next_vMerge = next_cell._tc.tcPr.xpath('./w:vMerge')
                                if next_vMerge:
                                    row_span += 1
                                else:
                                    break
                            else:
                                break
            
            return {
                "row_span": row_span,
                "col_span": grid_span
            }
            
        except Exception as e:
            logger.warning(f"Could not determine cell spans at ({row}, {col}): {str(e)}")
            return {
                "row_span": 1,
                "col_span": 1
            }

    def _is_merged_continuation(self, table: DocxTable, row: int, col: int) -> bool:
        """
        Check if cell is a continuation of a merged cell.
        
        Args:
            table: Word document table
            row: Row index
            col: Column index
            
        Returns:
            bool: True if cell is a merge continuation, False otherwise
        """
        try:
            cell = table.cell(row, col)
            tc = cell._tc
            if tc.tcPr is not None:
                vMerge = tc.tcPr.xpath('./w:vMerge')
                if vMerge:
                    # A cell is a continuation if it has vMerge without 'restart' value
                    return not (hasattr(vMerge[0], 'val') and vMerge[0].val == 'restart')
            return False
        except Exception as e:
            logger.warning(f"Could not check merge status for cell ({row}, {col}): {str(e)}")
            return False

    def _calculate_table_bbox(self, table: DocxTable) -> BoundingBox:
        """
        Calculate table bounding box coordinates.
        
        Args:
            table: Word document table
            
        Returns:
            BoundingBox containing table coordinates
        """
        try:
            # Get table width from table properties
            table_width = 0
            if hasattr(table, '_tbl') and table._tbl.tblPr is not None:
                tblW = table._tbl.tblPr.xpath('./w:tblW')
                if tblW and hasattr(tblW[0], 'w'):
                    # Convert width from twentieths of a point to inches
                    table_width = float(tblW[0].w) / 1440.0
            
            # If no width specified, estimate from content
            if table_width == 0:
                # Default to 6 inches if can't determine
                table_width = 6.0
            
            # Estimate height based on number of rows (approximate)
            row_height = 0.3  # Average row height in inches
            table_height = len(table.rows) * row_height
            
            # Use standard page width (8.5 inches) for relative positioning
            page_width = 8.5
            # Center the table horizontally if smaller than page width
            left_margin = (page_width - table_width) / 2 if table_width < page_width else 0
            
            return BoundingBox(
                x1=float(left_margin),
                y1=0.0,  # Relative to current position in document
                x2=float(left_margin + table_width),
                y2=float(table_height)
            )
            
        except Exception as e:
            logger.warning(f"Could not calculate exact table dimensions: {str(e)}")
            # Return default bounding box if calculation fails
            return BoundingBox(
                x1=0.0,
                y1=0.0,
                x2=6.0,  # Default width of 6 inches
                y2=float(len(table.rows) * 0.3)  # Approximate height
            )

    def _create_page_info(self, page_number: int, tables: List[DetectedTable]) -> PageTableInfo:
        """
        Create page information with tables.
        
        Args:
            page_number: Current page number
            tables: List of detected tables
            
        Returns:
            PageTableInfo containing page details and tables
        """
        return PageTableInfo(
            page_number=page_number,
            page_dimensions={
                "width": 8.5,  # Standard page width in inches
                "height": 11.0  # Standard page height in inches
            },
            tables=tables
        ) 