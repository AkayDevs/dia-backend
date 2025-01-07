from typing import Dict, Any, List
import docx
import logging
from pathlib import Path
from .base import BaseTableDetector, BaseTextExtractor
from app.core.config import settings

logger = logging.getLogger(__name__)

class WordTableDetector(BaseTableDetector):
    """Detector for tables in Word documents."""
    
    async def detect_tables(self, file_path: str, parameters: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Detect and extract tables from Word document."""
        try:
            # Convert to absolute path and check existence
            abs_path = Path(settings.UPLOAD_DIR) / file_path.replace("/uploads/", "")
            if not abs_path.exists():
                raise FileNotFoundError(f"Document not found at path: {abs_path}")
                
            doc = docx.Document(abs_path)
            min_row_count = parameters.get("min_row_count", 2)
            detect_headers = parameters.get("detect_headers", True)
            
            tables = []
            for table in doc.tables:
                if len(table.rows) < min_row_count:
                    continue
                    
                # Extract table data
                table_data = []
                headers = []
                
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    
                    if i == 0 and detect_headers:
                        headers = row_data
                    else:
                        table_data.append(row_data)
                
                # Create table info
                table_info = {
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": table_data,
                }
                
                if detect_headers and headers:
                    table_info["headers"] = headers
                
                tables.append(table_info)
                
            return tables
            
        except FileNotFoundError as e:
            logger.error(f"Document not found: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error detecting tables in Word document: {str(e)}")
            raise

class WordTextExtractor(BaseTextExtractor):
    """Extractor for text from Word documents."""
    
    async def extract_text(self, file_path: str, parameters: Dict[str, Any]) -> str:
        """Extract text from Word document."""
        try:
            # Convert to absolute path and check existence
            abs_path = Path(settings.UPLOAD_DIR) / file_path.replace("/uploads/", "")
            if not abs_path.exists():
                raise FileNotFoundError(f"Document not found at path: {abs_path}")
                
            doc = docx.Document(abs_path)
            extract_layout = parameters.get("extract_layout", True)
            detect_lists = parameters.get("detect_lists", True)
            
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    if extract_layout:
                        text_parts.append("")  # Preserve empty lines
                    continue
                
                # Handle lists if enabled
                if detect_lists and paragraph.style.name.startswith(("List", "Bullet")):
                    text_parts.append(f"• {paragraph.text}")
                else:
                    text_parts.append(paragraph.text)
            
            # Join with newlines if preserving layout, space otherwise
            separator = "\n" if extract_layout else " "
            return separator.join(text_parts)
            
        except FileNotFoundError as e:
            logger.error(f"Document not found: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from Word document: {str(e)}")
            raise 