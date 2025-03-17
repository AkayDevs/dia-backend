from typing import Dict, Any, List, Optional, Tuple
import os
import logging
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell
from PIL import Image
import math

from app.services.analysis.configs.base import BaseAlgorithm
from app.schemas.analysis.configs.algorithms import AlgorithmDefinitionBase, AlgorithmParameter, AlgorithmParameterValue
from app.enums.document import DocumentType
from app.schemas.analysis.results.table_structure import (
    TableStructureResult,
    PageTableStructureResult,
    TableStructure,
    Cell
)
from app.schemas.analysis.results.table_shared import BoundingBox, Confidence, PageInfo

logger = logging.getLogger(__name__)

class DocxTableStructureAlgorithm(BaseAlgorithm):
    """Table structure recognition for DOCX files using python-docx"""
    
    def get_info(self) -> AlgorithmDefinitionBase:
        return AlgorithmDefinitionBase(
            code="docx_structure",
            name="DOCX Table Structure Recognition",
            version="1.0.0",
            description="Table structure recognition for DOCX files using python-docx",
            supported_document_types=[DocumentType.DOCX],
            parameters=[
                AlgorithmParameter(
                    name="confidence_threshold",
                    description="Minimum confidence score for cell detection",
                    type="float",
                    required=False,
                    default=0.9,
                    constraints={
                        "min": 0.1,
                        "max": 1.0
                    }
                ),
                AlgorithmParameter(
                    name="header_row_count",
                    description="Number of rows to consider as header rows",
                    type="integer",
                    required=False,
                    default=1,
                    constraints={
                        "min": 0,
                        "max": 10
                    }
                )
            ],
            implementation_path="app.services.analysis.configs.definitions.table_analysis.algorithms.docx_structure.DocxTableStructureAlgorithm",
            is_active=True
        )

    def get_default_parameters(self) -> List[AlgorithmParameterValue]:
        """Get default parameters for the algorithm"""
        return [
            AlgorithmParameterValue(
                name="confidence_threshold",
                value=0.9
            ),
            AlgorithmParameterValue(
                name="header_row_count",
                value=1
            )
        ]
    
    async def validate_requirements(self) -> None:
        """Validate that all required dependencies are available"""
        try:
            from docx import Document
            from PIL import Image
        except ImportError as e:
            raise RuntimeError(f"Required dependency not found: {str(e)}")
    
    async def validate_input(self, input_data: Dict[str, Any]) -> None:
        """Validate input data format"""
        if "document_path" not in input_data:
            raise ValueError("Document path not provided in input data")
        if "tables" not in input_data:
            raise ValueError("Table detection results not provided in input data")
    
    def _map_docx_cell_to_pixel_coordinates(
        self, 
        cell: DocxCell, 
        table: DocxTable, 
        table_bbox: Dict[str, int],
        page_width: int,
        page_height: int,
        row_idx: int,
        col_idx: int
    ) -> Tuple[int, int, int, int]:
        """
        Map a DOCX cell to pixel coordinates based on its relative position in the table
        and the table's bounding box.
        """
        # Ensure all bbox values are integers
        x1 = int(table_bbox.get("x1", 0))
        y1 = int(table_bbox.get("y1", 0))
        x2 = int(table_bbox.get("x2", page_width))
        y2 = int(table_bbox.get("y2", page_height))
        
        # Get table dimensions
        table_width = max(1, x2 - x1)  # Ensure non-zero width
        table_height = max(1, y2 - y1)  # Ensure non-zero height
        
        # Calculate total rows and columns in the table
        total_rows = max(1, len(table.rows))  # Ensure non-zero rows
        total_cols = max(1, len(table.columns))  # Ensure non-zero columns
        
        # Calculate cell's relative position and size within the table
        # This is an approximation as we don't have exact measurements from python-docx
        cell_x1_rel = col_idx / total_cols
        cell_y1_rel = row_idx / total_rows
        
        # Calculate cell span
        row_span = 1
        col_span = 1
        
        # Try to get row and column spans if available
        try:
            # Check for vertical merge (row span)
            if hasattr(cell, "_element") and hasattr(cell._element, "vMerge"):
                span_value = cell._element.vMerge
                if span_value is not None:
                    row_span = int(span_value)
            
            # Check for horizontal merge (column span)
            if hasattr(cell, "_element") and hasattr(cell._element, "gridSpan"):
                span_value = cell._element.gridSpan
                if span_value is not None:
                    col_span = int(span_value)
                
            # Alternative way to check spans
            if hasattr(cell, "width") and cell.width:
                # If cell width is larger than average, it might be spanning
                avg_width = table_width / total_cols
                if cell.width > avg_width * 1.5:
                    col_span = max(col_span, 2)
        except Exception:
            # If we can't get spans, just use 1
            pass
        
        cell_x2_rel = (col_idx + col_span) / total_cols
        cell_y2_rel = (row_idx + row_span) / total_rows
        
        # Map to actual pixel coordinates
        cell_x1 = int(x1 + (cell_x1_rel * table_width))
        cell_y1 = int(y1 + (cell_y1_rel * table_height))
        cell_x2 = int(x1 + (cell_x2_rel * table_width))
        cell_y2 = int(y1 + (cell_y2_rel * table_height))
        
        # Ensure coordinates are within page boundaries
        cell_x1 = max(0, min(cell_x1, page_width))
        cell_y1 = max(0, min(cell_y1, page_height))
        cell_x2 = max(0, min(cell_x2, page_width))
        cell_y2 = max(0, min(cell_y2, page_height))
        
        return cell_x1, cell_y1, cell_x2, cell_y2
    
    def _process_docx_table(
        self,
        docx_table: DocxTable,
        table_bbox: Dict[str, int],
        page_width: int,
        page_height: int,
        header_row_count: int,
        confidence_score: float
    ) -> Tuple[List[Cell], int, int]:
        """Process a DOCX table to extract cells and table structure"""
        cells = []
        
        # Get table dimensions
        num_rows = len(docx_table.rows)
        num_cols = len(docx_table.columns)
        
        logger.info(f"Processing DOCX table with {num_rows} rows and {num_cols} columns")
        
        # Process each cell in the table
        for row_idx, row in enumerate(docx_table.rows):
            for col_idx, cell in enumerate(row.cells):
                try:
                    # Skip cells that are part of merged cells (not the top-left cell)
                    # This is a simplification as python-docx doesn't provide easy access to merged cells
                    
                    # Map cell to pixel coordinates
                    x1, y1, x2, y2 = self._map_docx_cell_to_pixel_coordinates(
                        cell, 
                        docx_table, 
                        table_bbox,
                        page_width,
                        page_height,
                        row_idx,
                        col_idx
                    )
                    
                    # Determine if this is a header cell
                    is_header = row_idx < header_row_count
                    
                    # Create cell object
                    cell_obj = Cell(
                        bbox=BoundingBox(
                            x1=x1,
                            y1=y1,
                            x2=x2,
                            y2=y2
                        ),
                        row_span=1,  # We don't have reliable row span info from python-docx
                        col_span=1,  # We don't have reliable col span info from python-docx
                        is_header=is_header,
                        confidence=Confidence(
                            score=confidence_score,
                            method="docx_structure"
                        )
                    )
                    cells.append(cell_obj)
                    
                    logger.debug(f"Added cell: row={row_idx}, col={col_idx}, bbox=({x1},{y1},{x2},{y2}), is_header={is_header}")
                except Exception as e:
                    logger.error(f"Error processing cell at row {row_idx}, col {col_idx}: {str(e)}")
                    # Continue with next cell instead of failing the entire table
                    continue
        
        return cells, num_rows, num_cols

    async def execute(
        self,
        document_path: str,
        parameters: Dict[str, Any],
        previous_results: Dict[str, Dict[str, Any]] = {}
    ) -> Dict[str, Any]:
        """Execute table structure recognition for DOCX files"""
        try:
            # Get parameters
            confidence_threshold = parameters.get("confidence_threshold", 0.9)
            header_row_count = parameters.get("header_row_count", 1)
            
            logger.info(f"Starting DOCX table structure recognition for {document_path}")
            
            # Extract document pages using document service
            from app.services.documents.document import extract_document_pages
            from app.enums.document import DocumentType

            # Ensure document is a DOCX file
            if not document_path.lower().endswith(('.docx', '.doc')):
                raise ValueError(f"Unsupported document type: {document_path}. Only DOCX files are supported.")

            # Extract user_id and document_id from path
            path_parts = document_path.split('/')
            if len(path_parts) < 3:
                raise ValueError(f"Invalid document path format: {document_path}")
            user_id = path_parts[0]
            document_id = path_parts[1]

            # Get document pages
            document_pages = await extract_document_pages(
                document_path=f"/uploads/{document_path}",
                document_type=DocumentType.DOCX,
                user_id=user_id,
                document_id=document_id
            )
            
            logger.info(f"Successfully extracted {len(document_pages.pages)} pages from document")
            
            # Get table detection results
            table_results = previous_results.get("table_analysis.table_detection", {})
            if not table_results:
                raise ValueError("No table detection results found in previous steps")
            
            # Open the DOCX file directly
            # Make sure to use the correct path format with leading slash
            docx_path = f"uploads/{document_path}"
            logger.info(f"Opening DOCX file at path: {docx_path}")
            
            try:
                docx = DocxDocument(docx_path)
                logger.info(f"Successfully opened DOCX file with {len(docx.tables)} tables")
            except Exception as e:
                logger.error(f"Error opening DOCX file: {str(e)}")
                raise ValueError(f"Failed to open DOCX file: {str(e)}")
            
            # Process each page
            final_results = []
            total_tables = 0
            
            for page_idx, page_result in enumerate(table_results.get("results", [])):
                try:
                    # Get corresponding document page
                    doc_page = next((p for p in document_pages.pages if p.page_number == page_idx + 1), None)
                    if not doc_page:
                        logger.warning(f"No document page found for page index {page_idx}")
                        continue

                    # Load image from the page's image_url to get dimensions
                    image_path = doc_page.image_url.lstrip('/')
                    logger.debug(f"Loading image from path: {image_path}")
                    image = Image.open(image_path)
                    page_width, page_height = image.size
                    logger.info(f"Page {page_idx+1} dimensions: {page_width}x{page_height}")

                    page_tables = []
                    page_info = page_result["page_info"]
                    
                    # In DOCX, we don't have a direct mapping between pages and tables
                    # We'll use the table detection results to identify tables on each page
                    for table_idx, table_loc in enumerate(page_result["tables"]):
                        try:
                            logger.info(f"Processing table {table_idx} on page {page_idx+1}")
                            # Get the corresponding table from the DOCX file
                            # This is an approximation as we don't have a direct mapping
                            if table_idx < len(docx.tables):
                                docx_table = docx.tables[table_idx]
                                
                                # Process the table
                                cells, num_rows, num_cols = self._process_docx_table(
                                    docx_table,
                                    table_loc["bbox"],
                                    page_width,
                                    page_height,
                                    header_row_count,
                                    confidence_threshold
                                )
                                
                                # Create table structure
                                table_structure = TableStructure(
                                    bbox=BoundingBox(**table_loc["bbox"]),
                                    cells=cells,
                                    num_rows=num_rows,
                                    num_cols=num_cols,
                                    confidence=Confidence(
                                        score=float(table_loc["confidence"]["score"]),
                                        method="docx_structure"
                                    )
                                )
                                
                                page_tables.append(table_structure)
                                total_tables += 1
                                logger.info(f"Successfully processed table {table_idx} with {len(cells)} cells")
                            else:
                                logger.warning(f"Table index {table_idx} exceeds available tables in DOCX ({len(docx.tables)})")
                        except Exception as e:
                            logger.error(f"Error processing table {table_idx} on page {page_idx+1}: {str(e)}")
                            # Continue with next table instead of failing the entire page
                            continue
                    
                    # Create page result
                    if page_tables:
                        page_result = PageTableStructureResult(
                            page_info=PageInfo(**page_info),
                            tables=page_tables,
                            processing_info={
                                "document_type": "DOCX",
                                "confidence_threshold": confidence_threshold,
                                "header_row_count": header_row_count
                            }
                        )
                        final_results.append(page_result)
                        logger.info(f"Added page result for page {page_idx+1} with {len(page_tables)} tables")
                except Exception as e:
                    logger.error(f"Error processing page {page_idx+1}: {str(e)}")
                    # Continue with next page instead of failing the entire document
                    continue
            
            # Create final result
            final_result = TableStructureResult(
                results=final_results,
                total_pages_processed=len(final_results),
                total_tables_processed=total_tables,
                metadata={
                    "algorithm": "docx_structure",
                    "version": "1.0.0",
                    "parameters": parameters,
                    "execution_info": {
                        "document_type": "DOCX"
                    }
                }
            )
            
            logger.info(f"DOCX table structure recognition completed successfully with {total_tables} tables across {len(final_results)} pages")
            return final_result.dict()
            
        except Exception as e:
            logger.error(f"DOCX table structure detection failed: {str(e)}")
            raise RuntimeError(f"DOCX table structure detection failed: {str(e)}")
    
    async def cleanup(self) -> None:
        """Cleanup any temporary resources"""
        pass 