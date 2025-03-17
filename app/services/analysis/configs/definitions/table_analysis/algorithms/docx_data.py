from typing import Dict, Any, List, Optional, Tuple
import os
import logging
import re
from datetime import datetime
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell
from PIL import Image

from app.services.analysis.configs.base import BaseAlgorithm
from app.schemas.analysis.configs.algorithms import AlgorithmDefinitionBase, AlgorithmParameter, AlgorithmParameterValue
from app.enums.document import DocumentType
from app.schemas.analysis.results.table_data import (
    TableDataResult,
    PageTableDataResult,
    TableData,
    CellContent
)
from app.schemas.analysis.results.table_shared import BoundingBox, Confidence, PageInfo

logger = logging.getLogger(__name__)

class DocxTableDataAlgorithm(BaseAlgorithm):
    """Table data extraction directly from DOCX files using python-docx"""
    
    def get_info(self) -> AlgorithmDefinitionBase:
        return AlgorithmDefinitionBase(
            code="docx_data",
            name="DOCX Table Data Extraction",
            version="1.0.0",
            description="Table data extraction directly from DOCX files using python-docx",
            supported_document_types=[DocumentType.DOCX],
            parameters=[
                AlgorithmParameter(
                    name="detect_data_types",
                    description="Whether to detect and normalize data types",
                    type="boolean",
                    required=False,
                    default=True
                ),
                AlgorithmParameter(
                    name="confidence_score",
                    description="Confidence score to assign to extracted data",
                    type="float",
                    required=False,
                    default=0.95,
                    constraints={
                        "min": 0.0,
                        "max": 1.0
                    }
                ),
                AlgorithmParameter(
                    name="extract_formatting",
                    description="Whether to extract text formatting information",
                    type="boolean",
                    required=False,
                    default=False
                )
            ],
            implementation_path="app.services.analysis.configs.definitions.table_analysis.algorithms.docx_data.DocxTableDataAlgorithm",
            is_active=True
        )

    def get_default_parameters(self) -> List[AlgorithmParameterValue]:
        """Get default parameters for the algorithm"""
        return [
            AlgorithmParameterValue(
                name="detect_data_types",
                value=True
            ),
            AlgorithmParameterValue(
                name="confidence_score",
                value=0.95
            ),
            AlgorithmParameterValue(
                name="extract_formatting",
                value=False
            )
        ]
    
    async def validate_requirements(self) -> None:
        """Validate that all required dependencies are available"""
        try:
            from docx import Document
            from PIL import Image
        except ImportError as e:
            raise RuntimeError(f"Required dependency not found: {str(e)}")
    
    async def validate_input(self, input_data: Dict[str, Any]) -> None:
        """Validate input data format"""
        if "document_path" not in input_data:
            raise ValueError("Document path not provided in input data")
        if "structures" not in input_data:
            raise ValueError("Table structure results not provided in input data")

    def _detect_data_type(self, text: str) -> tuple[str, Any]:
        """Detect data type and normalize value"""
        # Remove extra whitespace
        text = text.strip()
        if not text:
            return "empty", None

        # Check for numbers
        numeric_pattern = r'^-?\d*\.?\d+$'
        if re.match(numeric_pattern, text):
            try:
                value = float(text)
                return "number", value
            except ValueError:
                pass

        # Check for percentages
        percentage_pattern = r'^-?\d*\.?\d+\s*%$'
        if re.match(percentage_pattern, text):
            try:
                value = float(text.replace('%', '').strip()) / 100
                return "percentage", value
            except ValueError:
                pass

        # Check for dates (various formats)
        date_patterns = [
            ('%Y-%m-%d', r'^\d{4}-\d{2}-\d{2}$'),
            ('%d/%m/%Y', r'^\d{2}/\d{2}/\d{4}$'),
            ('%Y/%m/%d', r'^\d{4}/\d{2}/\d{2}$'),
            ('%d-%m-%Y', r'^\d{2}-\d{2}-\d{4}$'),
            ('%B %d, %Y', r'^[A-Za-z]+ \d{1,2}, \d{4}$')
        ]
        
        for date_format, pattern in date_patterns:
            if re.match(pattern, text):
                try:
                    value = datetime.strptime(text, date_format)
                    return "date", value.isoformat()
                except ValueError:
                    continue

        # Check for currency
        currency_pattern = r'^[$€£¥]?\s*-?\d*\.?\d+\s*[$€£¥]?$'
        if re.match(currency_pattern, text):
            try:
                # Extract numeric value
                value = float(re.findall(r'-?\d*\.?\d+', text)[0])
                return "currency", value
            except (ValueError, IndexError):
                pass

        # Default to text
        return "text", text

    def _extract_cell_content(
        self,
        cell: DocxCell,
        detect_types: bool,
        confidence_score: float,
        extract_formatting: bool
    ) -> CellContent:
        """Extract content from a single DOCX cell"""
        try:
            # Get text content from cell
            text = cell.text.strip() if cell.text else ""
            
            # Extract formatting information if requested
            formatting_info = {}
            if extract_formatting and hasattr(cell, "paragraphs"):
                try:
                    # Check if there are any paragraphs with runs
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        first_run = cell.paragraphs[0].runs[0]
                        if hasattr(first_run, "bold"):
                            formatting_info["bold"] = first_run.bold
                        if hasattr(first_run, "italic"):
                            formatting_info["italic"] = first_run.italic
                        if hasattr(first_run, "underline"):
                            formatting_info["underline"] = first_run.underline
                        if hasattr(first_run, "font") and hasattr(first_run.font, "size"):
                            formatting_info["font_size"] = first_run.font.size
                except Exception as e:
                    logger.warning(f"Error extracting formatting: {str(e)}")
            
            # Detect data type if enabled
            data_type = None
            normalized_value = None
            if detect_types and text:
                data_type, normalized_value = self._detect_data_type(text)
            
            return CellContent(
                text=text,
                confidence=Confidence(
                    score=float(confidence_score),
                    method="docx_direct_extraction"
                ),
                data_type=data_type,
                normalized_value=normalized_value,
                metadata=formatting_info if extract_formatting else None
            )
        except Exception as e:
            logger.error(f"Error extracting cell content: {str(e)}")
            # Return empty content on error
            return CellContent(
                text="",
                confidence=Confidence(
                    score=0.0,
                    method="docx_extraction_failed"
                ),
                data_type="empty",
                normalized_value=None
            )

    def _extract_table_data(
        self,
        docx_table: DocxTable,
        table_bbox: Dict[str, int],
        detect_types: bool,
        confidence_score: float,
        extract_formatting: bool
    ) -> List[List[CellContent]]:
        """Extract data from a DOCX table"""
        # Get table dimensions
        num_rows = len(docx_table.rows)
        num_cols = len(docx_table.columns)
        
        logger.info(f"Extracting data from DOCX table with {num_rows} rows and {num_cols} columns")
        
        # Initialize 2D array for table data
        table_data = []
        
        # Process each row
        for row_idx, row in enumerate(docx_table.rows):
            row_data = []
            
            # Process each cell in the row
            for col_idx, cell in enumerate(row.cells):
                try:
                    # Extract cell content
                    content = self._extract_cell_content(
                        cell,
                        detect_types,
                        confidence_score,
                        extract_formatting
                    )
                    row_data.append(content)
                    logger.debug(f"Extracted cell content at row {row_idx}, col {col_idx}: {content.text}")
                except Exception as e:
                    logger.error(f"Error processing cell at row {row_idx}, col {col_idx}: {str(e)}")
                    # Add empty content on error
                    row_data.append(CellContent(
                        text="",
                        confidence=Confidence(
                            score=0.0,
                            method="docx_extraction_failed"
                        ),
                        data_type="empty",
                        normalized_value=None
                    ))
            
            table_data.append(row_data)
        
        return table_data

    async def execute(
        self,
        document_path: str,
        parameters: Dict[str, Any],
        previous_results: Dict[str, Dict[str, Any]] = {}
    ) -> Dict[str, Any]:
        """Execute table data extraction from DOCX files"""
        try:
            # Get parameters
            detect_types = parameters.get("detect_data_types", True)
            confidence_score = parameters.get("confidence_score", 0.95)
            extract_formatting = parameters.get("extract_formatting", False)
            
            logger.info(f"Starting DOCX table data extraction for {document_path}")
            
            # Extract document pages using document service
            from app.services.documents.document import extract_document_pages
            from app.enums.document import DocumentType

            # Ensure document is a DOCX file
            if not document_path.lower().endswith(('.docx', '.doc')):
                raise ValueError(f"Unsupported document type: {document_path}. Only DOCX files are supported.")

            # Extract user_id and document_id from path
            path_parts = document_path.split('/')
            if len(path_parts) < 3:
                raise ValueError(f"Invalid document path format: {document_path}")
            user_id = path_parts[0]
            document_id = path_parts[1]

            # Get document pages
            document_pages = await extract_document_pages(
                document_path=f"/uploads/{document_path}",
                document_type=DocumentType.DOCX,
                user_id=user_id,
                document_id=document_id
            )
            
            logger.info(f"Successfully extracted {len(document_pages.pages)} pages from document")
            
            # Get structure results
            structure_results = previous_results.get("table_analysis.table_structure", {})
            if not structure_results:
                raise ValueError("No table structure results found in previous steps")
            
            # Open the DOCX file directly
            docx_path = f"uploads/{document_path}"
            logger.info(f"Opening DOCX file at path: {docx_path}")
            
            try:
                docx = DocxDocument(docx_path)
                logger.info(f"Successfully opened DOCX file with {len(docx.tables)} tables")
            except Exception as e:
                logger.error(f"Error opening DOCX file: {str(e)}")
                raise ValueError(f"Failed to open DOCX file: {str(e)}")
            
            # Process each page
            final_results = []
            total_tables = 0
            table_index = 0  # Global table index across all pages
            
            for page_result in structure_results.get("results", []):
                try:
                    # Get page info
                    page_info = page_result["page_info"]
                    page_tables = []
                    
                    # Process each table on the page
                    for table_structure in page_result["tables"]:
                        try:
                            # Check if we have enough tables in the document
                            if table_index < len(docx.tables):
                                docx_table = docx.tables[table_index]
                                
                                # Extract table data
                                table_data_cells = self._extract_table_data(
                                    docx_table,
                                    table_structure["bbox"],
                                    detect_types,
                                    confidence_score,
                                    extract_formatting
                                )
                                
                                # Create table data object
                                table = TableData(
                                    bbox=BoundingBox(**table_structure["bbox"]),
                                    cells=table_data_cells,
                                    confidence=Confidence(
                                        score=float(confidence_score),
                                        method="docx_direct_extraction"
                                    )
                                )
                                
                                page_tables.append(table)
                                total_tables += 1
                                logger.info(f"Successfully extracted data from table {table_index}")
                                
                                # Increment global table index
                                table_index += 1
                            else:
                                logger.warning(f"Table index {table_index} exceeds available tables in DOCX ({len(docx.tables)})")
                        except Exception as e:
                            logger.error(f"Error processing table {table_index}: {str(e)}")
                            # Continue with next table instead of failing the entire page
                            table_index += 1
                            continue
                    
                    # Create page result
                    if page_tables:
                        page_result = PageTableDataResult(
                            page_info=PageInfo(**page_info),
                            tables=page_tables,
                            processing_info={
                                "document_type": "DOCX",
                                "data_types_detected": detect_types,
                                "confidence_score": confidence_score,
                                "extract_formatting": extract_formatting
                            }
                        )
                        final_results.append(page_result)
                        logger.info(f"Added page result for page {page_info['page_number']} with {len(page_tables)} tables")
                except Exception as e:
                    logger.error(f"Error processing page {page_info['page_number']}: {str(e)}")
                    # Continue with next page instead of failing the entire document
                    continue
            
            # Create final result
            final_result = TableDataResult(
                results=final_results,
                total_pages_processed=len(final_results),
                total_tables_processed=total_tables,
                metadata={
                    "algorithm": "docx_data",
                    "version": "1.0.0",
                    "parameters": parameters,
                    "extraction_method": "python-docx"
                }
            )
            
            logger.info(f"DOCX table data extraction completed successfully with {total_tables} tables across {len(final_results)} pages")
            return final_result.dict()
            
        except Exception as e:
            logger.error(f"DOCX table data extraction failed: {str(e)}")
            raise RuntimeError(f"DOCX table data extraction failed: {str(e)}")
    
    async def cleanup(self) -> None:
        """Cleanup any temporary resources"""
        pass 