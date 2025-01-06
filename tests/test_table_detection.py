import pytest
import os
from pathlib import Path
import shutil
from PIL import Image, ImageDraw
import docx
import fitz  # PyMuPDF
import numpy as np
from datetime import datetime

from app.services.ml.table_detection import (
    ImageTableDetector,
    PDFTableDetector,
    WordTableDetector
)
from app.services.analysis import AnalysisOrchestrator
from app.schemas.analysis import AnalysisType, AnalysisStatus

# Test data directory
TEST_DATA_DIR = Path(__file__).parent / "test_data"
TEST_DATA_DIR.mkdir(exist_ok=True)

def create_test_image_with_table():
    """Create a test image containing a table."""
    # Create a white image
    img = Image.new('RGB', (800, 600), 'white')
    draw = ImageDraw.Draw(img)
    
    # Draw a simple table
    table_x = 100
    table_y = 100
    cell_width = 100
    cell_height = 50
    rows = 4
    cols = 3
    
    # Draw horizontal lines
    for i in range(rows + 1):
        y = table_y + i * cell_height
        draw.line([(table_x, y), (table_x + cols * cell_width, y)], fill='black', width=2)
    
    # Draw vertical lines
    for i in range(cols + 1):
        x = table_x + i * cell_width
        draw.line([(x, table_y), (x, table_y + rows * cell_height)], fill='black', width=2)
    
    # Save the image
    img_path = TEST_DATA_DIR / "test_table.png"
    img.save(img_path)
    return img_path

def create_test_pdf_with_table():
    """Create a test PDF containing a table."""
    doc = fitz.open()
    page = doc.new_page()
    
    # Draw table lines
    table_x = 100
    table_y = 100
    cell_width = 100
    cell_height = 50
    rows = 4
    cols = 3
    
    # Draw horizontal lines
    for i in range(rows + 1):
        y = table_y + i * cell_height
        page.draw_line(
            (table_x, y),
            (table_x + cols * cell_width, y),
            color=(0, 0, 0),
            width=1
        )
    
    # Draw vertical lines
    for i in range(cols + 1):
        x = table_x + i * cell_width
        page.draw_line(
            (x, table_y),
            (x, table_y + rows * cell_height),
            color=(0, 0, 0),
            width=1
        )
    
    # Add some text in cells
    page.insert_text((table_x + 10, table_y + 20), "Header 1")
    page.insert_text((table_x + cell_width + 10, table_y + 20), "Header 2")
    
    # Save the PDF
    pdf_path = TEST_DATA_DIR / "test_table.pdf"
    doc.save(pdf_path)
    doc.close()
    return pdf_path

def create_test_word_with_table():
    """Create a test Word document containing a table."""
    doc = docx.Document()
    
    # Add a heading
    doc.add_heading('Test Document with Table', 0)
    
    # Add some text
    doc.add_paragraph('This is a test document containing a table.')
    
    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Header 1'
    header_cells[1].text = 'Header 2'
    header_cells[2].text = 'Header 3'
    
    # Add some data
    data = [
        ['1,234.56', '2023-01-01', 'Text 1'],
        ['2,345.67', '2023-01-02', 'Text 2'],
        ['3,456.78', '2023-01-03', 'Text 3']
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        for j, value in enumerate(row_data):
            row[j].text = value
    
    # Save the document
    docx_path = TEST_DATA_DIR / "test_table.docx"
    doc.save(docx_path)
    return docx_path

@pytest.fixture(scope="module")
def test_files():
    """Create test files for all formats."""
    img_path = create_test_image_with_table()
    pdf_path = create_test_pdf_with_table()
    docx_path = create_test_word_with_table()
    
    yield {
        "image": img_path,
        "pdf": pdf_path,
        "word": docx_path
    }
    
    # Cleanup
    shutil.rmtree(TEST_DATA_DIR)

def test_image_table_detection(test_files):
    """Test table detection in images."""
    detector = ImageTableDetector()
    
    # Test file validation
    assert detector.validate_file(test_files["image"])
    
    # Test table detection
    tables = detector.detect_tables(
        test_files["image"],
        confidence_threshold=0.5,
        enhance_image=True
    )
    
    assert len(tables) > 0
    
    # Verify table structure
    table = tables[0]
    assert "coordinates" in table
    assert "confidence" in table
    assert "width" in table
    assert "height" in table
    assert table["width"] > 0
    assert table["height"] > 0

def test_pdf_table_detection(test_files):
    """Test table detection in PDFs."""
    detector = PDFTableDetector()
    
    # Test file validation
    assert detector.validate_file(test_files["pdf"])
    
    # Test table detection
    tables = detector.detect_tables(
        test_files["pdf"],
        confidence_threshold=0.5,
        use_ml_detection=True
    )
    
    assert len(tables) > 0
    
    # Verify table structure
    table = tables[0]
    assert "coordinates" in table
    assert "confidence" in table
    assert "page_number" in table
    assert table["page_number"] == 1

def test_word_table_detection(test_files):
    """Test table detection in Word documents."""
    detector = WordTableDetector()
    
    # Test file validation
    assert detector.validate_file(test_files["word"])
    
    # Test table detection with structure and data extraction
    tables = detector.detect_tables(
        test_files["word"],
        extract_structure=True,
        extract_data=True
    )
    
    assert len(tables) > 0
    
    # Verify table structure and data
    table = tables[0]
    assert "dimensions" in table
    assert "structure" in table
    assert "data" in table
    
    # Check structure
    structure = table["structure"]
    assert structure["has_header"]
    assert len(structure["header_rows"]) > 0
    assert structure["dimensions"]["rows"] == 4
    assert structure["dimensions"]["cols"] == 3
    
    # Check data types
    col_types = structure["column_types"]
    assert col_types[0]["type"] == "numeric"  # First column should be numeric
    assert col_types[1]["type"] == "date"     # Second column should be date
    assert col_types[2]["type"] == "text"     # Third column should be text
    
    # Check data extraction
    data = table["data"]
    assert len(data) == 3  # 3 data rows
    
    # Check header extraction
    headers = table["data"]["headers"]
    assert len(headers) == 3
    assert headers[0]["text"] == "Header 1"

@pytest.mark.asyncio
async def test_analysis_orchestration(test_files, mocker):
    """Test the complete analysis orchestration flow."""
    # Mock database session
    mock_db = mocker.MagicMock()
    
    # Create orchestrator
    orchestrator = AnalysisOrchestrator(mock_db)
    
    # Test for each file type
    for file_type, file_path in test_files.items():
        # Start analysis
        analysis = await orchestrator.start_analysis(
            document_id="test_doc_id",
            analysis_type=AnalysisType.TABLE_DETECTION,
            parameters={
                "confidence_threshold": 0.5,
                "extract_structure": True,
                "extract_data": True
            }
        )
        
        assert analysis is not None
        assert analysis.status == AnalysisStatus.PENDING
        
        # Process analysis
        await orchestrator._process_analysis(analysis.id)
        
        # Verify results
        mock_db.query.assert_called()
        
        # Check progress tracking
        progress_calls = [
            call for call in mock_db.method_calls 
            if "progress" in str(call)
        ]
        assert len(progress_calls) > 0

def test_error_handling(test_files):
    """Test error handling in table detection."""
    # Test with invalid files
    invalid_file = TEST_DATA_DIR / "invalid.xyz"
    invalid_file.touch()
    
    detectors = [
        ImageTableDetector(),
        PDFTableDetector(),
        WordTableDetector()
    ]
    
    for detector in detectors:
        # Should return False for invalid file
        assert not detector.validate_file(str(invalid_file))
        
        # Should raise exception when processing invalid file
        with pytest.raises(Exception):
            detector.detect_tables(str(invalid_file))
    
    # Test with corrupted files
    corrupted_file = TEST_DATA_DIR / "corrupted.pdf"
    with open(corrupted_file, "wb") as f:
        f.write(b"This is not a valid PDF file")
    
    pdf_detector = PDFTableDetector()
    assert not pdf_detector.validate_file(str(corrupted_file))
    
    # Cleanup
    invalid_file.unlink()
    corrupted_file.unlink()

if __name__ == "__main__":
    pytest.main([__file__, "-v"]) 